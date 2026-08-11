"""
TEA-Print Converter — núcleo de lógica (sem interface).

Usado pelo app.py (Streamlit). Mantém a mesma lógica pedagógica e de geração
de PDF já validada no notebook, mas devolvendo bytes em memória (BytesIO) em
vez de salvar direto em disco, para funcionar bem com download no navegador.
"""
import io
import re
import json
from pathlib import Path

import requests
import anthropic
import pdfplumber
import docx  # python-docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

try:
    import pytesseract
    OCR_DISPONIVEL = True
except ImportError:
    OCR_DISPONIVEL = False

# ---------------------------------------------------------------------------
# Configuração
# ---------------------------------------------------------------------------
CACHE_DIR = Path("cache_arasaac")
CACHE_DIR.mkdir(exist_ok=True)

ARASAAC_BASE = "https://api.arasaac.org/api"
ARASAAC_STATIC = "https://static.arasaac.org/pictograms"
ARASAAC_LANG = "pt"

MODELO_ANTHROPIC = "claude-sonnet-5"

NIVEIS = {
    "1": {"nome": "Nível 1", "sigla": "N1", "desc": "Requer apoio — linguagem simplificada, pistas visuais pontuais"},
    "2": {"nome": "Nível 2", "sigla": "N2", "desc": "Requer apoio substancial — chunking obrigatório, múltiplos pictogramas"},
    "3": {"nome": "Nível 3", "sigla": "N3", "desc": "Requer apoio muito substancial — máxima simplificação, CAA intensiva"},
}

CAMPOS_BNCC = [
    "O Eu, o Outro e o Nós",
    "Corpo, Gestos e Movimentos",
    "Traços, Sons, Cores e Formas",
    "Escuta, Fala, Pensamento e Imaginação",
    "Espaços, Tempos, Quantidades, Relações e Transformações",
]

SYSTEM_PROMPT = (
    "Você é especialista em adaptação pedagógica para alunos com TEA na Educação Infantil brasileira, "
    "alinhado à BNCC, Lei 13.146/2015 e Política Nacional de Educação Especial.\n"
    "Diretrizes obrigatórias:\n"
    "- Linguagem direta, sem metáforas ou ambiguidades\n"
    "- Instruções numeradas, curtas, uma ação por passo\n"
    "- Pictograma sugerido (palavra-chave em português, compatível com busca no ARASAAC) em cada passo\n"
    "- Nível 1: simplificar linguagem, pistas visuais pontuais\n"
    "- Nível 2: chunking obrigatório, múltiplos pictogramas\n"
    "- Nível 3: uma ação por frase, CAA intensiva\n"
    "Responda APENAS com JSON válido, sem texto extra, sem markdown."
)


# ---------------------------------------------------------------------------
# Leitura de arquivos
# ---------------------------------------------------------------------------
def ler_pdf_bytes(dados: bytes) -> str:
    texto = []
    with pdfplumber.open(io.BytesIO(dados)) as pdf:
        for pagina in pdf.pages:
            t = pagina.extract_text()
            if t:
                texto.append(t)
    return "\n".join(texto)


def ler_docx_bytes(dados: bytes) -> str:
    d = docx.Document(io.BytesIO(dados))
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


def ler_txt_bytes(dados: bytes) -> str:
    return dados.decode("utf-8", errors="ignore")


def ler_imagem_ocr_bytes(dados: bytes, idioma: str = "por") -> str:
    if not OCR_DISPONIVEL:
        raise RuntimeError(
            "OCR indisponível: o binário Tesseract não está instalado neste servidor. "
            "Cole o texto da atividade manualmente em vez de enviar imagem."
        )
    img = Image.open(io.BytesIO(dados))
    return pytesseract.image_to_string(img, lang=idioma)


def ler_arquivo_upload(nome_arquivo: str, dados: bytes) -> str:
    """Detecta o tipo pela extensão do nome do arquivo e extrai o texto."""
    sufixo = Path(nome_arquivo).suffix.lower()
    if sufixo == ".pdf":
        return ler_pdf_bytes(dados)
    elif sufixo == ".docx":
        return ler_docx_bytes(dados)
    elif sufixo == ".txt":
        return ler_txt_bytes(dados)
    elif sufixo in (".png", ".jpg", ".jpeg", ".webp"):
        return ler_imagem_ocr_bytes(dados)
    else:
        raise ValueError(f"Formato não suportado: {sufixo}")


# ---------------------------------------------------------------------------
# ARASAAC — https://arasaac.org/developers/api
# ---------------------------------------------------------------------------
def arasaac_buscar(palavra: str, idioma: str = ARASAAC_LANG):
    url = f"{ARASAAC_BASE}/pictograms/{idioma}/search/{requests.utils.quote(palavra)}"
    try:
        r = requests.get(url, timeout=8)
        r.raise_for_status()
        return r.json()
    except requests.RequestException:
        return []


def arasaac_baixar_imagem(pictograma_id, tamanho: int = 300):
    cache_path = CACHE_DIR / f"{pictograma_id}_{tamanho}.png"
    if cache_path.exists():
        return cache_path.read_bytes()
    url = f"{ARASAAC_STATIC}/{pictograma_id}/{pictograma_id}_{tamanho}.png"
    try:
        r = requests.get(url, timeout=8)
        r.raise_for_status()
        cache_path.write_bytes(r.content)
        return r.content
    except requests.RequestException:
        return None


def arasaac_melhor_pictograma(palavra_chave: str, tamanho: int = 300):
    if not palavra_chave:
        return None, None
    resultados = arasaac_buscar(palavra_chave)
    if not resultados:
        return None, None
    pictograma_id = resultados[0].get("_id")
    if pictograma_id is None:
        return None, None
    return pictograma_id, arasaac_baixar_imagem(pictograma_id, tamanho)


# ---------------------------------------------------------------------------
# Adaptação via API da Anthropic
# ---------------------------------------------------------------------------
def montar_prompt_usuario(nivel: str, campo: str, interesse: str) -> str:
    campo_txt = campo or "não especificado"
    interesse_txt = f"\nINTERESSE ESPECIAL: {interesse}" if interesse else ""
    exemplo_json = (
        '{"titulo":"título adaptado","campo_experiencia":"' + campo_txt + '",'
        '"nivel":"' + nivel + '","objetivo_pedagogico":"objetivo em frase única",'
        '"passos":[{"numero":1,"texto":"instrução direta","pictograma":"palavra-chave ARASAAC"}],'
        '"banco_palavras":[],"grade_caca_palavras":"",'
        '"zona_resposta":{"tipo":"linhas","descricao":"Espaço para resposta"},'
        '"orientacoes":{"objetivo":"objetivo com alinhamento BNCC",'
        '"como_apresentar":["passo 1","passo 2"],'
        '"dificuldades":[{"problema":"dificuldade","estrategia":"como contornar"}],'
        '"materiais":["material 1"],"caa":"sugestão CAA ou string vazia"}}'
    )
    return (
        f"Adapte esta atividade.\n\n"
        f"NÍVEL DSM-5: {nivel} — {NIVEIS[nivel]['desc']}\n"
        f"CAMPO BNCC: {campo_txt}{interesse_txt}\n\n"
        f"Retorne EXATAMENTE neste formato JSON:\n{exemplo_json}"
    )


class AdaptacaoIncompleta(Exception):
    """A resposta da IA foi cortada antes de terminar o JSON (limite de tokens atingido)."""
    def __init__(self, texto_bruto):
        self.texto_bruto = texto_bruto
        super().__init__(
            "A resposta da IA foi cortada antes de terminar (limite de tokens atingido). "
            "Tente novamente — se persistir, encurte um pouco o texto da atividade original."
        )


class AdaptacaoJsonInvalido(Exception):
    """A resposta da IA não é um JSON válido, mas não parece ter sido cortada por tokens."""
    def __init__(self, texto_bruto, erro_original):
        self.texto_bruto = texto_bruto
        self.erro_original = erro_original
        super().__init__(f"A resposta da IA não veio em um JSON válido: {erro_original}")


def _como_lista_str(valor):
    """Garante uma lista de strings, mesmo que a IA tenha retornado um formato diferente."""
    if isinstance(valor, list):
        return [str(x) for x in valor if x is not None and str(x).strip()]
    if isinstance(valor, str) and valor.strip():
        return [valor.strip()]
    return []


def _como_dict(valor):
    return valor if isinstance(valor, dict) else {}


def normalizar_adaptado(dados) -> dict:
    """Corrige pequenos desvios de formato que a IA às vezes comete (ex: um campo que
    deveria ser objeto vir como texto simples), para blindar a geração de PDF/DOCX
    contra AttributeError/TypeError."""
    dados = dados if isinstance(dados, dict) else {}
    out = {}

    out["titulo"] = str(dados.get("titulo") or "Atividade Adaptada")
    out["campo_experiencia"] = str(dados.get("campo_experiencia") or "")
    out["nivel"] = str(dados.get("nivel") or "1")
    out["objetivo_pedagogico"] = str(dados.get("objetivo_pedagogico") or "")

    passos_norm = []
    for i, p in enumerate(dados.get("passos") or [], start=1):
        if isinstance(p, dict):
            passos_norm.append({
                "numero": p.get("numero", i),
                "texto": str(p.get("texto") or ""),
                "pictograma": str(p.get("pictograma") or ""),
            })
        elif isinstance(p, str) and p.strip():
            passos_norm.append({"numero": i, "texto": p.strip(), "pictograma": ""})
    out["passos"] = passos_norm

    out["banco_palavras"] = _como_lista_str(dados.get("banco_palavras"))
    out["grade_caca_palavras"] = str(dados.get("grade_caca_palavras") or "")

    zona = dados.get("zona_resposta")
    if isinstance(zona, dict):
        out["zona_resposta"] = {
            "tipo": str(zona.get("tipo") or "linhas"),
            "descricao": str(zona.get("descricao") or "Espaço para resposta"),
        }
    elif isinstance(zona, str) and zona.strip():
        out["zona_resposta"] = {"tipo": "linhas", "descricao": zona.strip()}
    else:
        out["zona_resposta"] = {"tipo": "linhas", "descricao": "Espaço para resposta"}

    o = _como_dict(dados.get("orientacoes"))
    dificuldades_norm = []
    for d in (o.get("dificuldades") or []):
        if isinstance(d, dict):
            dificuldades_norm.append({
                "problema": str(d.get("problema") or ""),
                "estrategia": str(d.get("estrategia") or ""),
            })
        elif isinstance(d, str) and d.strip():
            dificuldades_norm.append({"problema": d.strip(), "estrategia": ""})
    out["orientacoes"] = {
        "objetivo": str(o.get("objetivo") or ""),
        "como_apresentar": _como_lista_str(o.get("como_apresentar")),
        "dificuldades": dificuldades_norm,
        "materiais": _como_lista_str(o.get("materiais")),
        "caa": str(o.get("caa") or ""),
    }
    return out


def adaptar_atividade(api_key: str, texto_original: str, nivel: str = "1",
                       campo: str = "", interesse: str = "", modelo: str = None,
                       max_tokens: int = 4096) -> dict:
    client = anthropic.Anthropic(api_key=api_key)
    modelo = modelo or MODELO_ANTHROPIC
    prompt_usuario = "ATIVIDADE ORIGINAL:\n" + texto_original + "\n\n" + montar_prompt_usuario(nivel, campo, interesse)
    resposta = client.messages.create(
        model=modelo,
        max_tokens=max_tokens,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt_usuario}],
    )
    bruto = "".join(bloco.text for bloco in resposta.content if bloco.type == "text")
    bruto = re.sub(r"```json|```", "", bruto).strip()

    try:
        dados = json.loads(bruto)
    except json.JSONDecodeError as e:
        if getattr(resposta, "stop_reason", None) == "max_tokens":
            raise AdaptacaoIncompleta(bruto) from e
        raise AdaptacaoJsonInvalido(bruto, e) from e

    return normalizar_adaptado(dados)


# ---------------------------------------------------------------------------
# Geração de PDF (retorna bytes)
# ---------------------------------------------------------------------------
def _quebra_texto(c, texto, largura_max, fonte="Helvetica", tamanho=11):
    c.setFont(fonte, tamanho)
    palavras = str(texto).split()
    linhas, linha = [], ""
    for p in palavras:
        teste = (linha + " " + p).strip()
        if c.stringWidth(teste, fonte, tamanho) <= largura_max:
            linha = teste
        else:
            if linha:
                linhas.append(linha)
            linha = p
    if linha:
        linhas.append(linha)
    return linhas or [""]


def gerar_pdf_bytes(adaptado: dict, titulo_atividade: str, nome_aluno: str = "",
                     nivel: str = "1", usar_pictogramas: bool = True) -> bytes:
    nivel_info = NIVEIS[nivel]
    M = 20 * mm
    W, H = A4
    CW = W - 2 * M

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    def rodape_atribuicao():
        if usar_pictogramas:
            c.setFont("Helvetica", 6)
            c.setFillColorRGB(0.6, 0.6, 0.6)
            c.drawString(M, 10 * mm, "Pictogramas: ARASAAC (arasaac.org) — CC BY-NC-SA 4.0 — Governo de Aragón")
            c.setFillColorRGB(0, 0, 0)

    def checar_quebra(y, necessario):
        if y - necessario < M:
            rodape_atribuicao()
            c.showPage()
            return H - M
        return y

    def titulo_secao(y, texto):
        c.setFont("Helvetica-Bold", 10)
        c.setFillColorRGB(0.2, 0.2, 0.2)
        c.drawString(M, y, texto.upper())
        c.setLineWidth(0.4)
        c.line(M, y - 2 * mm, W - M, y - 2 * mm)
        c.setFillColorRGB(0, 0, 0)
        return y - 6 * mm

    # ---------- PÁGINA 1: ATIVIDADE DO ALUNO ----------
    y = H - M
    c.setFont("Helvetica-Bold", 8)
    c.setFillColorRGB(0.47, 0.47, 0.47)
    c.drawString(M, y, "TEA · ATIVIDADE ADAPTADA")
    c.setFillColorRGB(0, 0, 0)

    c.setFont("Helvetica", 12)
    c.drawString(M, y - 10 * mm, f"Nome: {nome_aluno or '_' * 35}")
    c.setFont("Helvetica", 11)
    c.drawRightString(W - M, y - 4 * mm, "Data: ____/____/______")
    c.drawRightString(W - M, y - 10 * mm, "Turma: ___________")

    c.rect(W - M - 15 * mm, y - 16 * mm, 15 * mm, 6 * mm)
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(W - M - 7.5 * mm, y - 14 * mm, nivel_info["sigla"])

    y -= 20 * mm
    c.setLineWidth(0.7)
    c.line(M, y, W - M, y)

    y -= 10 * mm
    for lt in _quebra_texto(c, titulo_atividade, CW, "Helvetica-Bold", 19):
        c.setFont("Helvetica-Bold", 19)
        c.drawCentredString(W / 2, y, lt)
        y -= 8 * mm
    y -= 4 * mm

    for passo in adaptado.get("passos", []):
        texto = passo.get("texto", "")
        picto_kw = passo.get("pictograma", "")
        linhas_passo = _quebra_texto(c, texto, CW - 36 * mm, "Helvetica-Bold", 13)
        altura_linha = max(23 * mm, len(linhas_passo) * 6 * mm + 6 * mm)
        y = checar_quebra(y, altura_linha)
        y_topo = y

        c.setFillColorRGB(0.13, 0.13, 0.13)
        c.circle(M + 4.5 * mm, y_topo - 4.5 * mm, 4.5 * mm, fill=1)
        c.setFillColorRGB(1, 1, 1)
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(M + 4.5 * mm, y_topo - 6.3 * mm, str(passo.get("numero", "")))
        c.setFillColorRGB(0, 0, 0)

        px = M + 12 * mm
        caixa = 21 * mm
        imagem_bytes = None
        if usar_pictogramas and picto_kw:
            _, imagem_bytes = arasaac_melhor_pictograma(picto_kw)
        c.setDash(1.5, 1.5)
        c.setStrokeColorRGB(0.7, 0.7, 0.7)
        c.rect(px, y_topo - caixa, caixa, caixa)
        c.setDash()
        c.setStrokeColorRGB(0, 0, 0)
        if imagem_bytes:
            try:
                img = ImageReader(io.BytesIO(imagem_bytes))
                c.drawImage(img, px + 1 * mm, y_topo - caixa + 1 * mm,
                            width=caixa - 2 * mm, height=caixa - 2 * mm,
                            preserveAspectRatio=True, mask="auto")
            except Exception:
                imagem_bytes = None
        if not imagem_bytes:
            c.setFont("Helvetica", 7)
            c.setFillColorRGB(0.55, 0.55, 0.55)
            c.drawCentredString(px + caixa / 2, y_topo - 10 * mm, (picto_kw or "")[:16])
            c.drawCentredString(px + caixa / 2, y_topo - 17 * mm, "pictograma")
            c.setFillColorRGB(0, 0, 0)

        c.setFont("Helvetica-Bold", 13)
        yt = y_topo - 6 * mm
        for lp in linhas_passo:
            c.drawString(px + 25 * mm, yt, lp)
            yt -= 6 * mm

        y = y_topo - altura_linha

    banco = adaptado.get("banco_palavras") or []
    if banco:
        y = checar_quebra(y, 25 * mm)
        y -= 3 * mm
        c.setFont("Helvetica-Bold", 9)
        c.drawString(M + 3 * mm, y, "BANCO DE PALAVRAS")
        y -= 8 * mm
        wx = M + 4 * mm
        for palavra in banco:
            c.setFont("Helvetica-Bold", 13)
            largura = c.stringWidth(palavra, "Helvetica-Bold", 13) + 8 * mm
            if wx + largura > W - M - 3 * mm:
                wx = M + 4 * mm
                y -= 11 * mm
            c.rect(wx, y - 5.5 * mm, largura, 8 * mm)
            c.drawString(wx + 4 * mm, y, palavra)
            wx += largura + 3 * mm
        y -= 10 * mm

    grade = (adaptado.get("grade_caca_palavras") or "").strip()
    if grade:
        linhas_grade = [l for l in grade.split("\n") if l.strip()]
        altura_grade = len(linhas_grade) * 7 * mm + 8 * mm
        y = checar_quebra(y, altura_grade)
        y -= 3 * mm
        c.setFont("Courier-Bold", 12)
        gy = y - 7 * mm
        for gl in linhas_grade:
            c.drawString(M + 4 * mm, gy, gl)
            gy -= 7 * mm
        y -= altura_grade

    zona = adaptado.get("zona_resposta")
    if zona:
        y = checar_quebra(y, 46 * mm)
        y -= 5 * mm
        c.setFont("Helvetica", 9)
        c.setFillColorRGB(0.4, 0.4, 0.4)
        c.drawString(M, y + 2 * mm, zona.get("descricao", "Espaço para resposta"))
        c.setFillColorRGB(0, 0, 0)
        altura_zona = 38 * mm
        c.setLineWidth(0.7)
        c.rect(M, y - altura_zona, CW, altura_zona)
        c.setLineWidth(0.2)
        c.setStrokeColorRGB(0.82, 0.82, 0.82)
        for i in range(1, 4):
            c.line(M + 3 * mm, y - i * 9.5 * mm, M + CW - 3 * mm, y - i * 9.5 * mm)
        c.setStrokeColorRGB(0, 0, 0)

    rodape_atribuicao()
    c.showPage()

    # ---------- PÁGINA 2: ORIENTAÇÕES DO PROFESSOR ----------
    y = H - M
    o = adaptado.get("orientacoes", {})

    c.setFillColorRGB(0.93, 0.93, 0.93)
    c.rect(M, y - 18 * mm, CW, 18 * mm, fill=1, stroke=0)
    c.setFillColorRGB(0, 0, 0)
    c.setFont("Helvetica-Bold", 15)
    c.drawString(M + 5 * mm, y - 8 * mm, "Orientações Pedagógicas")
    c.setFont("Helvetica", 10)
    c.setFillColorRGB(0.35, 0.35, 0.35)
    sub = f"{titulo_atividade} · {nivel_info['nome']}"
    if adaptado.get("campo_experiencia"):
        sub += f" · {adaptado['campo_experiencia']}"
    c.drawString(M + 5 * mm, y - 14 * mm, sub[:110])
    c.setFillColorRGB(0, 0, 0)
    y -= 23 * mm

    objetivo = o.get("objetivo") or adaptado.get("objetivo_pedagogico", "")
    if objetivo:
        y = checar_quebra(y, 30 * mm)
        y = titulo_secao(y, "Objetivo")
        for l in _quebra_texto(c, objetivo, CW, "Helvetica", 11):
            y = checar_quebra(y, 6 * mm)
            c.setFont("Helvetica", 11)
            c.drawString(M, y, l)
            y -= 5.5 * mm
        y -= 4 * mm

    como = o.get("como_apresentar") or []
    if como:
        y = checar_quebra(y, 30 * mm)
        y = titulo_secao(y, "Como apresentar ao aluno")
        for i, passo in enumerate(como, 1):
            linhas = _quebra_texto(c, passo, CW - 9 * mm, "Helvetica", 11)
            y = checar_quebra(y, len(linhas) * 5.5 * mm + 2 * mm)
            c.setFont("Helvetica-Bold", 11)
            c.drawString(M, y, f"{i}.")
            c.setFont("Helvetica", 11)
            for l in linhas:
                c.drawString(M + 8 * mm, y, l)
                y -= 5.5 * mm
            y -= 2 * mm
        y -= 3 * mm

    dificuldades = o.get("dificuldades") or []
    if dificuldades:
        y = checar_quebra(y, 40 * mm)
        y = titulo_secao(y, "Dificuldades e estratégias")
        cw2 = (CW - 3 * mm) / 2
        c.setFillColorRGB(0.93, 0.93, 0.93)
        c.rect(M, y - 7 * mm, cw2, 7 * mm, fill=1, stroke=0)
        c.rect(M + cw2 + 3 * mm, y - 7 * mm, cw2, 7 * mm, fill=1, stroke=0)
        c.setFillColorRGB(0, 0, 0)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(M + 3 * mm, y - 5 * mm, "Dificuldade")
        c.drawString(M + cw2 + 6 * mm, y - 5 * mm, "Estratégia")
        y -= 7 * mm
        for item in dificuldades:
            pl = _quebra_texto(c, item.get("problema", ""), cw2 - 5 * mm, "Helvetica", 10)
            el = _quebra_texto(c, item.get("estrategia", ""), cw2 - 5 * mm, "Helvetica", 10)
            altura = max(len(pl), len(el)) * 5 * mm + 6 * mm
            y = checar_quebra(y, altura + 2 * mm)
            c.setLineWidth(0.3)
            c.rect(M, y - altura, cw2, altura)
            c.rect(M + cw2 + 3 * mm, y - altura, cw2, altura)
            c.setFont("Helvetica", 10)
            yy = y - 5 * mm
            for l in pl:
                c.drawString(M + 3 * mm, yy, l)
                yy -= 5 * mm
            yy = y - 5 * mm
            for l in el:
                c.drawString(M + cw2 + 6 * mm, yy, l)
                yy -= 5 * mm
            y -= altura
        y -= 4 * mm

    materiais = o.get("materiais") or []
    if materiais:
        y = checar_quebra(y, 30 * mm)
        y = titulo_secao(y, "Materiais e recursos")
        for m in materiais:
            linhas = _quebra_texto(c, m, CW - 7 * mm, "Helvetica", 11)
            y = checar_quebra(y, len(linhas) * 5.5 * mm + 2 * mm)
            c.setFont("Helvetica-Bold", 11)
            c.drawString(M, y, "•")
            c.setFont("Helvetica", 11)
            for l in linhas:
                c.drawString(M + 6 * mm, y, l)
                y -= 5.5 * mm
            y -= 2 * mm
        y -= 3 * mm

    caa = o.get("caa", "")
    if caa:
        y = checar_quebra(y, 30 * mm)
        y = titulo_secao(y, "Uso de CAA")
        for l in _quebra_texto(c, caa, CW, "Helvetica", 11):
            y = checar_quebra(y, 6 * mm)
            c.setFont("Helvetica", 11)
            c.drawString(M, y, l)
            y -= 5.5 * mm

    rodape_atribuicao()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Geração de DOCX (editável pelo professor) — retorna bytes
# ---------------------------------------------------------------------------
def _config_pagina_a4(document, margem_cm=2.0):
    for secao in document.sections:
        secao.page_height = Cm(29.7)
        secao.page_width = Cm(21.0)
        secao.top_margin = Cm(margem_cm)
        secao.bottom_margin = Cm(margem_cm)
        secao.left_margin = Cm(margem_cm)
        secao.right_margin = Cm(margem_cm)


def _estilo_base(document):
    estilo = document.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(12)


def _celula_texto(cell, texto, tamanho=11, negrito=False, cor=None, alinhamento=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alinhamento is not None:
        p.alignment = alinhamento
    run = p.add_run(str(texto))
    run.font.size = Pt(tamanho)
    run.font.bold = negrito
    if cor:
        run.font.color.rgb = RGBColor(*cor)
    return p


def _titulo_secao_docx(document, texto):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto.upper())
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # linha inferior simples via borda de parágrafo
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "999999",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def gerar_docx_bytes(adaptado: dict, titulo_atividade: str, nome_aluno: str = "",
                      nivel: str = "1", usar_pictogramas: bool = True) -> bytes:
    """Gera um .docx editável: página 1 = atividade do aluno, página 2 = orientações do professor."""
    nivel_info = NIVEIS[nivel]
    document = Document()
    _config_pagina_a4(document)
    _estilo_base(document)

    # ---------- CABEÇALHO ----------
    p = document.add_paragraph()
    r = p.add_run("TEA · ATIVIDADE ADAPTADA")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    tabela_cab = document.add_table(rows=1, cols=2)
    tabela_cab.autofit = True
    _celula_texto(tabela_cab.cell(0, 0), f"Nome: {nome_aluno or '_' * 35}", tamanho=12)
    _celula_texto(tabela_cab.cell(0, 1),
                   f"Data: ____/____/______      Turma: ___________      Nível: {nivel_info['sigla']}",
                   tamanho=11, alinhamento=WD_ALIGN_PARAGRAPH.RIGHT)

    # ---------- TÍTULO ----------
    titulo_p = document.add_heading(level=1)
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_p.add_run(titulo_atividade)
    titulo_run.font.size = Pt(20)
    titulo_run.font.color.rgb = RGBColor(0x08, 0x50, 0x41)

    document.add_paragraph()

    # ---------- PASSOS (tabela editável: nº | pictograma | instrução) ----------
    passos = adaptado.get("passos", [])
    if passos:
        tabela = document.add_table(rows=0, cols=3)
        tabela.style = "Table Grid"
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabela.columns[0].width = Cm(1.2)
        tabela.columns[1].width = Cm(3.0)
        tabela.columns[2].width = Cm(12.8)

        for passo in passos:
            linha = tabela.add_row()
            _celula_texto(linha.cells[0], str(passo.get("numero", "")), tamanho=13, negrito=True,
                          alinhamento=WD_ALIGN_PARAGRAPH.CENTER)

            cel_picto = linha.cells[1]
            cel_picto.text = ""
            picto_kw = passo.get("pictograma", "")
            imagem_bytes = None
            if usar_pictogramas and picto_kw:
                _, imagem_bytes = arasaac_melhor_pictograma(picto_kw)
            par_picto = cel_picto.paragraphs[0]
            par_picto.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if imagem_bytes:
                try:
                    par_picto.add_run().add_picture(io.BytesIO(imagem_bytes), width=Cm(2.2))
                except Exception:
                    imagem_bytes = None
            if not imagem_bytes:
                run_p = par_picto.add_run(picto_kw or "")
                run_p.font.size = Pt(8)
                run_p.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            _celula_texto(linha.cells[2], passo.get("texto", ""), tamanho=13, negrito=True)

        document.add_paragraph()

    # ---------- BANCO DE PALAVRAS ----------
    banco = adaptado.get("banco_palavras") or []
    if banco:
        p = document.add_paragraph()
        r = p.add_run("BANCO DE PALAVRAS")
        r.font.bold = True
        r.font.size = Pt(10)
        p2 = document.add_paragraph()
        r2 = p2.add_run("   ".join(banco))
        r2.font.bold = True
        r2.font.size = Pt(14)
        document.add_paragraph()

    # ---------- GRADE DE CAÇA-PALAVRAS ----------
    grade = (adaptado.get("grade_caca_palavras") or "").strip()
    if grade:
        for linha_txt in grade.split("\n"):
            if not linha_txt.strip():
                continue
            p = document.add_paragraph()
            r = p.add_run(linha_txt)
            r.font.name = "Courier New"
            r.font.size = Pt(13)
            r.font.bold = True
        document.add_paragraph()

    # ---------- ZONA DE RESPOSTA ----------
    zona = adaptado.get("zona_resposta")
    if zona:
        p = document.add_paragraph()
        r = p.add_run(zona.get("descricao", "Espaço para resposta"))
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        tabela_zona = document.add_table(rows=1, cols=1)
        tabela_zona.style = "Table Grid"
        cel = tabela_zona.cell(0, 0)
        cel.text = "\n" * 6  # espaço em branco para o aluno escrever/desenhar

    # ---------- QUEBRA DE PÁGINA → ORIENTAÇÕES DO PROFESSOR ----------
    document.add_page_break()

    o = adaptado.get("orientacoes", {})
    titulo_prof = document.add_heading(level=1)
    r = titulo_prof.add_run("Orientações Pedagógicas")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x08, 0x50, 0x41)

    sub = f"{titulo_atividade} · {nivel_info['nome']}"
    if adaptado.get("campo_experiencia"):
        sub += f" · {adaptado['campo_experiencia']}"
    p_sub = document.add_paragraph()
    r_sub = p_sub.add_run(sub)
    r_sub.font.size = Pt(10)
    r_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    objetivo = o.get("objetivo") or adaptado.get("objetivo_pedagogico", "")
    if objetivo:
        _titulo_secao_docx(document, "Objetivo")
        document.add_paragraph(objetivo)

    como = o.get("como_apresentar") or []
    if como:
        _titulo_secao_docx(document, "Como apresentar ao aluno")
        for passo in como:
            document.add_paragraph(passo, style="List Number")

    dificuldades = o.get("dificuldades") or []
    if dificuldades:
        _titulo_secao_docx(document, "Dificuldades e estratégias")
        tabela_d = document.add_table(rows=1, cols=2)
        tabela_d.style = "Table Grid"
        _celula_texto(tabela_d.cell(0, 0), "Dificuldade", tamanho=10, negrito=True)
        _celula_texto(tabela_d.cell(0, 1), "Estratégia", tamanho=10, negrito=True)
        for item in dificuldades:
            linha = tabela_d.add_row()
            _celula_texto(linha.cells[0], item.get("problema", ""), tamanho=10)
            _celula_texto(linha.cells[1], item.get("estrategia", ""), tamanho=10)

    materiais = o.get("materiais") or []
    if materiais:
        _titulo_secao_docx(document, "Materiais e recursos")
        for m in materiais:
            document.add_paragraph(m, style="List Bullet")

    caa = o.get("caa", "")
    if caa:
        _titulo_secao_docx(document, "Uso de CAA")
        document.add_paragraph(caa)

    if usar_pictogramas:
        document.add_paragraph()
        p_rodape = document.add_paragraph()
        r_rodape = p_rodape.add_run(
            "Pictogramas: ARASAAC (arasaac.org) — CC BY-NC-SA 4.0 — Governo de Aragón"
        )
        r_rodape.font.size = Pt(7)
        r_rodape.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
